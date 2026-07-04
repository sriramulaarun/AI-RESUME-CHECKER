import re
import pdfplumber
import docx

# Initialize spacy model lazily if installed
HAS_SPACY = False
_nlp = None

try:
    import spacy
    from spacy.matcher import PhraseMatcher
    HAS_SPACY = True
except ImportError:
    print("spaCy is not installed. Using pure-Python NLP fallback parser.")

def get_nlp():
    global _nlp, HAS_SPACY
    if not HAS_SPACY:
        return None
        
    if _nlp is None:
        try:
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Download model if not present on system
            try:
                from spacy.cli import download
                download("en_core_web_sm")
                _nlp = spacy.load("en_core_web_sm")
            except Exception as e:
                print(f"Error downloading spaCy model: {e}")
                HAS_SPACY = False
                return None
    return _nlp

# Comprehensive list of technology and industry skills
SKILLS_LIST = [
    # Languages
    "Python", "Java", "JavaScript", "JS", "TypeScript", "TS", "C", "C++", "C#", "Ruby", "PHP", 
    "Swift", "Go", "Golang", "Rust", "Kotlin", "Scala", "R", "SQL", "HTML", "CSS", "Bash", "Shell",
    
    # Frontend
    "React", "ReactJS", "React.js", "Angular", "AngularJS", "Vue", "VueJS", "Vue.js", 
    "Svelte", "Next.js", "Nuxt.js", "jQuery", "Bootstrap", "Tailwind", "TailwindCSS",
    
    # Backend
    "Django", "Flask", "FastAPI", "Node.js", "NodeJS", "Express.js", "Express", 
    "Spring Boot", "Spring", "Ruby on Rails", "Rails", "ASP.NET", "Laravel",
    
    # Cloud / DevOps
    "AWS", "Amazon Web Services", "GCP", "Google Cloud", "Google Cloud Platform", 
    "Azure", "Microsoft Azure", "Docker", "Kubernetes", "K8s", "CI/CD", "Git", "GitHub", 
    "GitLab", "Jenkins", "Terraform", "Ansible", "Linux", "Unix",
    
    # Databases
    "PostgreSQL", "Postgres", "MySQL", "MongoDB", "Redis", "SQLite", "Cassandra", 
    "DynamoDB", "Oracle", "Firebase",
    
    # AI / ML / Data Science
    "Machine Learning", "Deep Learning", "Artificial Intelligence", "AI", "ML", 
    "Natural Language Processing", "NLP", "Computer Vision", "spaCy", "NLTK", 
    "TensorFlow", "PyTorch", "Keras", "Scikit-Learn", "Sklearn", "Pandas", "NumPy", 
    "Tableau", "Power BI", "Data Analysis", "Data Science",
    
    # PM & Methodologies
    "Agile", "Scrum", "Project Management", "Product Management", "System Design", 
    "Software Engineering", "SDLC", "Jira", "Confluence",
    
    # Soft skills
    "Leadership", "Communication", "Teamwork", "Problem Solving", "Time Management", "Critical Thinking"
]

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file including paragraphs and tables."""
    text = []
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        # Dedup text from merged cells if they appear multiple times
                        cell_text = cell.text.strip()
                        if not text or text[-1] != cell_text:
                            text.append(cell_text)
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return "\n".join(text)

def extract_email(text):
    """Extract email using regular expression."""
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    match = re.search(email_pattern, text)
    return match.group(0) if match else None

def extract_phone(text):
    """Extract phone number using regular expression."""
    # Matches formats like +1 234 567 8900, (123) 456-7890, 123-456-7890, 1234567890
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    match = re.search(phone_pattern, text)
    return match.group(0) if match else None

def extract_name(text):
    """Extract candidate name using spaCy NER (if available) or layout heuristics."""
    # Split text into non-empty lines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return "Unknown Candidate"
        
    nlp = get_nlp()
    if nlp is not None:
        # Usually the name is in the first 3 lines of the resume. Let's look for PERSON entities there.
        name_candidates = []
        for line in lines[:3]:
            # Skip lines that look like emails, URLs, or addresses
            if '@' in line or 'http' in line or 'www.' in line or re.search(r'\d{5}', line):
                continue
            doc = nlp(line)
            for ent in doc.ents:
                if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                    if not any(word in ent.text.lower() for word in ["curriculum", "vitae", "resume", "page", "phone", "email"]):
                        name_candidates.append(ent.text)
        if name_candidates:
            return name_candidates[0]
            
    # Pure-Python Fallback heuristic:
    # Check the first 3 lines for a line of 2-4 capitalized words that looks like a name
    for line in lines[:3]:
        # Filter out contact details
        if '@' in line or 'http' in line or 'www.' in line or re.search(r'\d{5}', line) or '|' in line or ':' in line:
            continue
        # Check if the line consists of capitalized words only
        words = line.split()
        if 2 <= len(words) <= 4:
            # Must start with uppercase letters and consist of alphabetical characters
            if all(w[0].isupper() and w.isalpha() for w in words):
                # Ensure no common section labels
                if not any(w.lower() in ["curriculum", "vitae", "resume", "summary", "profile", "contact"] for w in words):
                    return line
                    
    # Ultimate fallback: pick the first line if it's short
    first_line = lines[0]
    if 2 <= len(first_line.split()) <= 4 and not any(char in first_line for char in ['@', ':', '/', '\\', '|']):
        if not any(word in first_line.lower() for word in ["curriculum", "vitae", "resume"]):
            return first_line
        
    return "Unknown Candidate"

def extract_skills(text):
    """Extract skills using spaCy PhraseMatcher (if available) or pure-Python regex fallback."""
    nlp = get_nlp()
    if nlp is not None:
        try:
            doc = nlp(text)
            matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
            patterns = [nlp.make_doc(skill) for skill in SKILLS_LIST]
            matcher.add("SKILLS", patterns)
            matches = matcher(doc)
            
            skills = set()
            for match_id, start, end in matches:
                span = doc[start:end]
                matched_text = span.text
                for s in SKILLS_LIST:
                    if s.lower() == matched_text.lower():
                        skills.add(s)
                        break
                else:
                    skills.add(matched_text.title())
            return sorted(list(skills))
        except Exception as e:
            print(f"Error using spaCy matcher: {e}. Falling back to Python matcher.")

    # Pure-Python Fallback:
    # Search the text case-insensitively using regex word boundaries
    text_lower = text.lower()
    skills = set()
    for skill in SKILLS_LIST:
        escaped_skill = re.escape(skill.lower())
        
        # Determine appropriate word boundary regex
        # If the skill starts/ends with alphanumeric characters, apply \b boundaries
        pattern = escaped_skill
        if skill[0].isalnum():
            pattern = r'\b' + pattern
        if skill[-1].isalnum():
            pattern = pattern + r'\b'
            
        if re.search(pattern, text_lower):
            skills.add(skill)
            
    return sorted(list(skills))

def extract_education(text):
    """Extract education details based on section search and keyword indexing."""
    edu_keywords = [
        "education", "academic qualification", "qualifications", 
        "academic background", "academics", "educational background",
        "academic profile", "degrees"
    ]
    
    degree_keywords = [
        "bachelor", "master", "ph.d", "doctor of", "b.s", "m.s", "b.tech", "m.tech", 
        "b.e", "m.e", "mba", "b.a", "m.a", "bsc", "msc", "phd", "undergraduate", "graduate",
        "university", "college", "school", "institute", "polytechnic"
    ]
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    education_lines = []
    
    # Find indices where education section may start
    edu_section_idx = -1
    for i, line in enumerate(lines):
        if any(keyword == line.lower().strip(': ') for keyword in edu_keywords) or (
            any(keyword in line.lower() for keyword in edu_keywords) and len(line) < 30
        ):
            edu_section_idx = i
            break
            
    if edu_section_idx != -1:
        # Extract lines following section header until we see another common section header
        section_headers = ["experience", "work", "history", "employment", "project", "skill", "certification", "award", "interest", "language", "profile", "summary"]
        for j in range(edu_section_idx + 1, min(edu_section_idx + 15, len(lines))):
            line = lines[j]
            # Check if this line is likely the start of another section
            if (any(h == line.lower().strip(': ') for h in section_headers) or (
                any(h in line.lower() for h in section_headers) and len(line) < 25
            )):
                break
            education_lines.append(line)
            
    # Fallback/Supplemental: Scan the entire document for lines containing degree/university info
    if len(education_lines) < 2:
        temp_lines = []
        for line in lines:
            if any(keyword in line.lower() for keyword in degree_keywords):
                temp_lines.append(line)
        # Deduplicate and append lines that aren't already included
        for l in temp_lines:
            if l not in education_lines:
                education_lines.append(l)
                
    if not education_lines:
        return "Not Specified"
        
    return "\n".join(education_lines)

def parse_resume(file_path, filename):
    """Core function to extract fields from a resume."""
    ext = filename.rsplit('.', 1)[-1].lower()
    
    if ext == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        text = extract_text_from_docx(file_path)
    else:
        # Text/fallback parsing
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            text = ""
            
    if not text.strip():
        return {
            "name": "Unknown Candidate",
            "email": None,
            "phone": None,
            "skills": [],
            "education": "Not Specified",
            "raw_text": ""
        }
        
    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    skills = extract_skills(text)
    education = extract_education(text)
    
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "education": education,
        "raw_text": text
    }
